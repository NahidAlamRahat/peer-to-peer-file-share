import subprocess, sys

# Install python-docx if not installed
subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_image_centered(path, width=Inches(2.5), caption=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

# ── Title ──────────────────────────────────────────
title = doc.add_heading('অ্যাসাইনমেন্ট', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('বিষয়: Peer-to-Peer (P2P) ফাইল শেয়ারিং অ্যাপ্লিকেশন')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(13)

doc.add_paragraph()
info = doc.add_paragraph()
info.add_run('শিক্ষার্থীর নাম: ').bold = True
info.add_run('[তোমার নাম]\n')
info.add_run('কোর্স: ').bold = True
info.add_run('[তোমার কোর্স]\n')
info.add_run('বিভাগ: ').bold = True
info.add_run('[তোমার বিভাগ]\n')
info.add_run('তারিখ: ').bold = True
info.add_run('২২ এপ্রিল, ২০২৬')

doc.add_page_break()

# ── Section 1 ──────────────────────────────────────
heading('১. ভূমিকা (Introduction)')
para('আধুনিক ডিজিটাল যুগে ফাইল শেয়ারিং একটি অত্যন্ত গুরুত্বপূর্ণ প্রয়োজনীয়তা। বর্তমানে Google Drive, WeTransfer বা Telegram-এর মতো সার্ভার-নির্ভর প্ল্যাটফর্মগুলো ব্যবহার করে ফাইল শেয়ার করতে হয়। এতে ফাইলটি প্রথমে একটি কেন্দ্রীয় সার্ভারে আপলোড হয়, তারপর অন্যজন ডাউনলোড করেন। এই পদ্ধতিতে গোপনীয়তা লঙ্ঘনের ঝুঁকি থাকে, ফাইলের সাইজ সীমাবদ্ধতা থাকে এবং সার্ভার খরচ বেশি হয়।')
para('এই সমস্যাগুলো সমাধান করতে আমি একটি Peer-to-Peer (P2P) ফাইল শেয়ারিং অ্যাপ্লিকেশন তৈরি করেছি। এই অ্যাপে দুইটি ডিভাইস সরাসরি একে অপরের সাথে সংযুক্ত হয়ে ফাইল আদান-প্রদান করতে পারে — কোনো মধ্যস্থতাকারী সার্ভারে ফাইল সংরক্ষণ ছাড়াই।')

# ── Section 2 ──────────────────────────────────────
heading('২. প্রজেক্টের উদ্দেশ্য (Objective)')
for item in [
    'ইন্টারনেটের মাধ্যমে দুটি ডিভাইসের মধ্যে সরাসরি ফাইল ট্রান্সফার করা',
    'কোনো ফাইল সাইজ সীমা ছাড়াই যেকোনো ফাইল পাঠানো',
    'ফাইল ট্রান্সফারে End-to-End Encryption নিশ্চিত করা',
    'QR কোড বা লিংক শেয়ারের মাধ্যমে সহজে সংযোগ স্থাপন করা',
    'Android, iOS, Web, Windows, macOS, Linux — সকল প্ল্যাটফর্মে কাজ করা',
]:
    doc.add_paragraph(item, style='List Bullet')

# ── Section 3 ──────────────────────────────────────
heading('৩. ব্যবহৃত প্রযুক্তি ও ভাষাসমূহ (Technologies & Languages Used)')
heading('৩.১ প্রোগ্রামিং ভাষা', level=2)
para('• Dart — সম্পূর্ণ অ্যাপের মূল কোড লেখা হয়েছে Dart ভাষায়')
para('• JavaScript/Node.js — Signaling Server তৈরিতে ব্যবহার করা হয়েছে')

heading('৩.২ Framework ও Platform', level=2)
para('Flutter — Google-এর তৈরি এই ফ্রেমওয়ার্ক একটি মাত্র কোডবেস থেকে Android, iOS, Web, Windows, macOS এবং Linux সব প্ল্যাটফর্মে অ্যাপ চালাতে পারে।')

heading('৩.৩ মূল লাইব্রেরিসমূহ', level=2)
libs = [
    ('flutter_webrtc ^1.3.1', 'WebRTC প্রোটোকল ব্যবহার করে P2P সংযোগ তৈরি'),
    ('flutter_bloc ^9.1.1', 'State Management (BLoC Pattern)'),
    ('web_socket_channel ^3.0.3', 'Signaling Server-এর সাথে WebSocket সংযোগ'),
    ('file_picker ^10.3.10', 'ডিভাইস থেকে ফাইল নির্বাচন'),
    ('qr_flutter ^4.1.0', 'QR Code তৈরি'),
    ('permission_handler ^12.0.1', 'Android/iOS পারমিশন ম্যানেজমেন্ট'),
    ('flutter_local_notifications ^21.0.0', 'ট্রান্সফার প্রোগ্রেস নোটিফিকেশন'),
    ('path_provider ^2.1.5', 'ফাইল সেভ করার পাথ নির্ধারণ'),
    ('shared_preferences ^2.5.4', 'Settings সংরক্ষণ'),
    ('wakelock_plus ^1.5.1', 'ট্রান্সফারের সময় স্ক্রিন বন্ধ না হওয়া নিশ্চিত করা'),
    ('get_it ^9.2.1', 'Dependency Injection'),
    ('uuid ^4.5.3', 'প্রতিটি ফাইলের জন্য ইউনিক আইডি তৈরি'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'লাইব্রেরি'
hdr[1].text = 'কাজ'
for lib, use in libs:
    row = table.add_row().cells
    row[0].text = lib
    row[1].text = use

heading('৩.৪ Deployment Platform', level=2)
para('• Render — Node.js Signaling Server হোস্ট করা হয়েছে')
para('• Vercel — Web অ্যাপ হোস্ট করা হয়েছে')

# ── Section 4 ──────────────────────────────────────
heading('৪. অ্যাপের আর্কিটেকচার (Architecture)')
para('এই প্রজেক্টে Clean Architecture এবং BLoC (Business Logic Component) Pattern অনুসরণ করা হয়েছে।')
heading('৪.১ লেয়ার বিভাজন', level=2)
for item in [
    'core/ — Theme, DI, Services (সকল লেয়ারে শেয়ার হওয়া কোড)',
    'data/ → datasources/ — WebRTC Client, Signaling Service',
    'data/ → repositories/ — File Transfer Repository Implementation',
    'domain/ → entities/ — ShareFile, FileChunkInfo, PeerSession',
    'domain/ → repositories/ — Abstract Repository Interfaces',
    'presentation/ → blocs/ — ConnectionBloc, TransferBloc',
    'presentation/ → screens/ — HomeScreen, ShareLinkScreen, ReceiveScreen, TransferScreen',
    'presentation/ → widgets/ — Reusable UI components',
]:
    doc.add_paragraph(item, style='List Bullet')

heading('৪.২ কিভাবে কাজ করে', level=2)
para('ধাপ ১ — Signaling (সংযোগ স্থাপন):', bold=True)
for item in [
    'Sender অ্যাপ খুলে "Send Files" বাটনে ক্লিক করে',
    'Signaling Server একটি ইউনিক Session ID তৈরি করে',
    'Sender সেই Session ID-এর QR Code বা লিংক Receiver-কে পাঠায়',
    'Receiver লিংক স্ক্যান করলে WebRTC Handshake হয় (Offer → Answer → ICE Candidates)',
]:
    doc.add_paragraph(item, style='List Number')

para('ধাপ ২ — File Transfer (ডেটা ট্রান্সফার):', bold=True)
for item in [
    'WebRTC DataChannel খোলে',
    'ফাইলটি ১৬KB চাংকে ভাগ করে পাঠানো হয়',
    'প্রতিটি চাংক Receiver-এ জোড়া লাগানো হয়',
    'সম্পূর্ণ ফাইল পেলে অটোমেটিক সেভ হয়',
    'Signaling Server-এ কোনো ফাইল ডেটা যায় না',
]:
    doc.add_paragraph(item, style='List Number')

# ── Section 5 ──────────────────────────────────────
heading('৫. মূল সমস্যাসমূহ ও সমাধান (Problems & Solutions)')

problems = [
    ('সমস্যা ১: NAT Traversal সমস্যা',
     'বাস্তব জীবনে বেশিরভাগ ডিভাইস NAT-এর পিছনে থাকে। এই কারণে দুটি ডিভাইস সরাসরি একে অপরকে খুঁজে পেতে পারে না।',
     'একাধিক STUN Server ব্যবহার করা হয়েছে (stun.l.google.com)। STUN Server ডিভাইসের Public IP ও Port বের করে দেয়।'),
    ('সমস্যা ২: Signaling Server Connection বিচ্ছিন্ন হওয়া',
     'Render-এ হোস্ট করা Signaling Server দীর্ঘ নিষ্ক্রিয়তায় Sleep Mode-এ চলে যায়। পুনরায় সংযোগ করতে গেলে 404 Error দেখায়।',
     'Exponential Backoff Reconnection (৩s, ৬s, ১২s, ২৪s, ৩০s) এবং প্রতি ২০ সেকেন্ডে Heartbeat Ping/Pong যোগ করা হয়েছে।'),
    ('সমস্যা ৩: বড় ফাইল ট্রান্সফারে Memory Overflow',
     'বড় ফাইল (১GB+) একসাথে মেমোরিতে লোড করলে অ্যাপ ক্র্যাশ করে।',
     'ফাইল ১৬KB Chunk-এ কেটে পাঠানো হয়। Buffer পূর্ণ হলে পাঠানো থামিয়ে অপেক্ষা করে। UI আপডেট ১০০ms-এ Throttle করা হয়েছে।'),
    ('সমস্যা ৪: Transfer State UI রিসেট না হওয়া',
     'ট্রান্সফার সম্পন্ন হওয়ার পরেও "Active Transfer" ব্যানার স্ক্রিনে থেকে যাচ্ছিল।',
     'BLoC-এ ResetTransferEvent যোগ করা হয়েছে। ব্যবহারকারী ✕ বাটনে ক্লিক করে ক্লিয়ার করতে পারেন।'),
    ('সমস্যা ৫: Multi-platform File Saving',
     'Android, iOS, Web এবং Desktop-এ ফাইল সেভ করার পদ্ধতি সম্পূর্ণ ভিন্ন।',
     'Conditional Import ব্যবহার করা হয়েছে। Web-এ Browser download API এবং Mobile/Desktop-এ path_provider ব্যবহার করা হয়।'),
]

for title_text, problem, solution in problems:
    heading(title_text, level=2)
    p = doc.add_paragraph()
    p.add_run('সমস্যা: ').bold = True
    p.add_run(problem)
    s = doc.add_paragraph()
    s.add_run('সমাধান: ').bold = True
    s.add_run(solution)

# ── Section 6 ──────────────────────────────────────
heading('৬. নিরাপত্তা বিশ্লেষণ (Security Analysis)')
heading('৬.১ End-to-End Encryption (E2EE)', level=2)
para('WebRTC DTLS-SRTP Encryption ব্যবহার করা হয়েছে। এটি স্বয়ংক্রিয়ভাবে সমস্ত DataChannel ট্র্যাফিক এনক্রিপ্ট করে। ফলে মাঝপথে কেউ ডেটা চুরি করতে পারে না।')
para('• DTLS: Connection স্থাপনের সময় Certificate যাচাই করে')
para('• SRTP: ডেটা ট্রান্সফারের সময় এনক্রিপশন নিশ্চিত করে')

heading('৬.২ No Server Storage', level=2)
para('Signaling Server শুধুমাত্র সংযোগ স্থাপনে সাহায্য করে। কোনো ফাইলের ডেটা সার্ভারে যায় না।')

heading('৬.৩ Session-based Access Control', level=2)
para('প্রতিটি ট্রান্সফারের জন্য একটি ইউনিক Session ID তৈরি হয়। এই আইডি না জানলে কেউ সংযোগ করতে পারে না।')

# ── Section 7 ──────────────────────────────────────
heading('৭. প্ল্যাটফর্ম সমর্থন (Platform Support)')
platforms = [
    ('Android', '✅ সম্পূর্ণ সমর্থিত'),
    ('iOS', '✅ সমর্থিত'),
    ('Web (Browser)', '✅ সমর্থিত (Vercel-এ Deployed)'),
    ('Windows', '✅ সমর্থিত'),
    ('macOS', '✅ সমর্থিত'),
    ('Linux', '✅ সমর্থিত'),
]
pt = doc.add_table(rows=1, cols=2)
pt.style = 'Table Grid'
ph = pt.rows[0].cells
ph[0].text = 'প্ল্যাটফর্ম'
ph[1].text = 'সমর্থন'
for pl, sup in platforms:
    row = pt.add_row().cells
    row[0].text = pl
    row[1].text = sup

# ── Section 8 — Screenshots ────────────────────────
heading('৮. অ্যাপের স্ক্রিনসমূহ (App Screens)')
para('নিচে অ্যাপ্লিকেশনের বিভিন্ন স্ক্রিনের বাস্তব Screenshot দেওয়া হলো:')
doc.add_paragraph()

screens = [
    (r'D:\rahat\1.jpeg', 'চিত্র ১: Home Screen — Send Files ও Receive Files বাটন'),
    (r'D:\rahat\2.jpeg', 'চিত্র ২: Share File Screen — QR Code ও Session Code (191393)'),
    (r'D:\rahat\5.jpeg', 'চিত্র ৩: Receive File Screen — ফাইল গ্রহণের জন্য Download বাটন'),
    (r'D:\rahat\4.jpeg', 'চিত্র ৪: Live Transfer Screen — রিয়েল-টাইম Progress ও Speed'),
    (r'D:\rahat\3.jpeg', 'চিত্র ৫: Transfer Complete Screen — সফল ট্রান্সফারের পর সবুজ চেকমার্ক'),
]

for img_path, caption in screens:
    add_image_centered(img_path, width=Inches(2.5), caption=caption)
    doc.add_paragraph()

# ── Section 9 ──────────────────────────────────────
heading('৯. প্রজেক্টের সীমাবদ্ধতা (Limitations)')
for item in [
    'TURN Server অনুপস্থিত: Symmetric NAT পরিবেশে সংযোগ ব্যর্থ হতে পারে।',
    'একটি Session-এ একজন Receiver: বর্তমানে একটি লিংক শুধু একজনের সাথে শেয়ার করা যায়।',
    'Render Free Tier: Signaling Server বিনামূল্যে হওয়ায় দীর্ঘ নিষ্ক্রিয়তায় Sleep Mode-এ যায়।',
]:
    doc.add_paragraph(item, style='List Number')

# ── Section 10 ─────────────────────────────────────
heading('১০. উপসংহার (Conclusion)')
para('এই প্রজেক্টটি WebRTC প্রযুক্তি ব্যবহার করে একটি সম্পূর্ণ Peer-to-Peer ফাইল শেয়ারিং সিস্টেম তৈরির সফল উদ্যোগ। এটি প্রমাণ করে যে আধুনিক ওয়েব প্রযুক্তি ব্যবহার করে কোনো কেন্দ্রীয় সার্ভার ছাড়াই নিরাপদ, দ্রুত এবং সীমাহীন ফাইল ট্রান্সফার সম্ভব।')
para('Flutter ফ্রেমওয়ার্ক ব্যবহার করে একটি কোডবেস থেকে ৬টি প্ল্যাটফর্মে অ্যাপ চালানো, Clean Architecture-এর মাধ্যমে কোড পরিষ্কার রাখা এবং বাস্তব সমস্যাগুলো (NAT Traversal, Reconnection, Buffer Management) সমাধান করা — এই প্রজেক্টের মূল অর্জন।')
para('ভবিষ্যতে TURN Server যোগ, Multi-peer support এবং File Transfer Password Protection যোগ করলে এটি একটি পূর্ণাঙ্গ প্রোডাকশন-রেডি অ্যাপ্লিকেশন হয়ে উঠবে।')

doc.add_paragraph()
ref = doc.add_paragraph('GitHub Repo: NahidAlamRahat/peer-to-peer-file-share')
ref.runs[0].italic = True

# Save
out = r'D:\rahat\assignment.docx'
doc.save(out)
print(f'✅ Word file saved: {out}')
